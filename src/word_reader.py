"""
Word文書・テキストファイル読み込みモジュール
"""

from docx import Document
import io
from typing import Optional


def read_word_file(file_bytes: bytes) -> dict:
    """
    Wordファイルを読み込み、構造化データを返す
    """
    doc = Document(io.BytesIO(file_bytes))
    
    paragraphs = []
    full_text_parts = []
    image_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                'text': text,
                'style': para.style.name if para.style else 'Normal',
                'is_bold': any(run.bold for run in para.runs if run.bold),
            })
            full_text_parts.append(text)
        
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                image_count += 1
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    
    image_count = max(1, image_count // 2) if image_count > 0 else 0
    
    return {
        'full_text': '\n\n'.join(full_text_parts),
        'paragraphs': paragraphs,
        'has_images': image_count > 0,
        'image_count': image_count,
    }


def extract_text_only(file_bytes: bytes, file_type: str = "docx") -> str:
    """
    ファイルからテキストのみを抽出
    
    Args:
        file_bytes: ファイルのバイトデータ
        file_type: ファイル形式 ("docx" または "txt")
        
    Returns:
        str: 抽出されたテキスト
    """
    if file_type == "txt":
        # テキストファイルの場合
        try:
            # UTF-8で試す
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Shift-JISで試す（日本語Windows）
                return file_bytes.decode('shift_jis')
            except UnicodeDecodeError:
                # CP932で試す
                return file_bytes.decode('cp932', errors='ignore')
    else:
        # Wordファイルの場合
        doc = Document(io.BytesIO(file_bytes))
        
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        
        return '\n\n'.join(text_parts)
